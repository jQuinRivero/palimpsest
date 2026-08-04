from __future__ import annotations

from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document as WordDocument
from docx.enum.style import WD_STYLE_TYPE

from app.models.document import BlockKind, Document
from app.services.ingestion.base import DocumentSource, SourceProbe
from app.services.ingestion.docx import DocxParser


def _docx_bytes(paragraphs: list[tuple[str, str | None]]) -> bytes:
    document = WordDocument()
    for text, style in paragraphs:
        paragraph = document.add_paragraph(text)
        if style is not None:
            paragraph.style = style
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _parse(data: bytes, filename: str = "witness.docx") -> Document:
    return DocxParser().parse(
        DocumentSource(
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            size_bytes=len(data),
            data=data,
        )
    )


def _replace_zip_part(data: bytes, replacements: dict[str, bytes]) -> bytes:
    source = BytesIO(data)
    target = BytesIO()
    with ZipFile(source) as input_zip, ZipFile(target, "w", ZIP_DEFLATED) as output_zip:
        for item in input_zip.infolist():
            output_zip.writestr(
                item, replacements.get(item.filename, input_zip.read(item.filename))
            )
        for name, content in replacements.items():
            if name not in input_zip.namelist():
                output_zip.writestr(name, content)
    return target.getvalue()


def test_docx_can_parse_only_wordprocessingml_zip() -> None:
    data = _docx_bytes([("Body", None)])
    xlsx_like = BytesIO()
    with ZipFile(xlsx_like, "w", ZIP_DEFLATED) as archive:
        archive.writestr("xl/workbook.xml", "<workbook />")

    assert DocxParser.can_parse(
        SourceProbe(
            filename="renamed.txt", media_type="text/plain", magic_bytes=data, size_bytes=len(data)
        )
    )
    assert not DocxParser.can_parse(
        SourceProbe(
            filename="fake.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            magic_bytes=xlsx_like.getvalue(),
            size_bytes=len(xlsx_like.getvalue()),
        )
    )
    assert not DocxParser.can_parse(
        SourceProbe(filename="fake.docx", media_type=None, magic_bytes=b"%PDF-", size_bytes=5)
    )


def test_docx_can_parse_declared_docx_without_magic_bytes() -> None:
    assert DocxParser.can_parse(
        SourceProbe(filename="witness.docx", media_type=None, magic_bytes=b"", size_bytes=0)
    )
    assert DocxParser.can_parse(
        SourceProbe(
            filename=None,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            magic_bytes=b"",
            size_bytes=0,
        )
    )


def test_block_kind_assignment_for_docx_styles() -> None:
    data = _docx_bytes(
        [
            ("Chapter", "Heading 1"),
            ("Quoted", "Quote"),
            ("Intense", "Intense Quote"),
            ("Listed", "List Bullet"),
            ("Body", None),
        ]
    )
    document = _parse(data)

    assert [block.kind for block in document.blocks] == [
        BlockKind.HEADING,
        BlockKind.QUOTE,
        BlockKind.QUOTE,
        BlockKind.LIST_ITEM,
        BlockKind.PARAGRAPH,
    ]
    assert [block.style for block in document.blocks] == [
        "Heading 1",
        "Quote",
        "Intense Quote",
        "List Bullet",
        "Normal",
    ]


def test_capabilities_match_docx_heading_behavior() -> None:
    capabilities = DocxParser.capabilities()
    document = _parse(_docx_bytes([("Chapter", "Heading 2"), ("Body", None)]))

    assert capabilities.preserves_headings is True
    assert [block.kind for block in document.blocks] == [
        BlockKind.HEADING,
        BlockKind.PARAGRAPH,
    ]


def test_localized_heading_style_name_is_preserved_as_heading() -> None:
    document_fixture = WordDocument()
    localized_heading = document_fixture.styles.add_style("Título 1", WD_STYLE_TYPE.PARAGRAPH)
    document_fixture.add_paragraph("Capítulo", style=localized_heading)
    buffer = BytesIO()
    document_fixture.save(buffer)

    document = _parse(buffer.getvalue())

    assert document.blocks[0].kind is BlockKind.HEADING
    assert document.blocks[0].style == "Título 1"
    assert document.warnings == []


def test_empty_paragraphs_are_skipped() -> None:
    data = _docx_bytes([("", None), ("   ", None), ("Body", None)])
    document = _parse(data)

    assert [block.text for block in document.blocks] == ["Body"]


def test_empty_docx_produces_empty_document() -> None:
    document = _parse(_docx_bytes([]))

    assert document.blocks == []
    assert document.full_text() == ""
    assert document.metadata.block_count == 0


def test_malformed_or_non_docx_zip_is_rejected() -> None:
    fake_zip = BytesIO()
    with ZipFile(fake_zip, "w", ZIP_DEFLATED) as archive:
        archive.writestr("xl/workbook.xml", "<workbook />")

    with pytest.raises(ValueError, match=r"word/document\.xml"):
        _parse(fake_zip.getvalue())


def test_docx_warnings_for_ignored_non_body_content() -> None:
    data = _docx_bytes([("Body", None)])
    with ZipFile(BytesIO(data)) as archive:
        document_xml = archive.read("word/document.xml")
    document_xml = document_xml.replace(
        b"<w:t>Body</w:t>",
        b"<w:ins><w:r><w:t>Inserted</w:t></w:r></w:ins><w:txbxContent><w:p><w:r><w:t>Box</w:t></w:r></w:p></w:txbxContent><w:t>Body</w:t>",
    )
    data = _replace_zip_part(
        data,
        {
            "word/document.xml": document_xml,
            "word/comments.xml": b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" />',
            "word/footnotes.xml": b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" />',
            "word/endnotes.xml": b'<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" />',
        },
    )

    document = _parse(data)

    assert set(warning.code for warning in document.warnings) == {
        "TRACKED_CHANGES_IGNORED",
        "COMMENTS_IGNORED",
        "FOOTNOTES_IGNORED",
        "ENDNOTES_IGNORED",
        "TEXT_BOXES_IGNORED",
    }


def test_unknown_non_default_style_warns_and_preserves_style() -> None:
    document_fixture = WordDocument()
    custom_style = document_fixture.styles.add_style("Custom Aside", WD_STYLE_TYPE.PARAGRAPH)
    document_fixture.add_paragraph("Aside", style=custom_style)
    buffer = BytesIO()
    document_fixture.save(buffer)

    document = _parse(buffer.getvalue())

    assert document.blocks[0].kind is BlockKind.PARAGRAPH
    assert document.blocks[0].style == "Custom Aside"
    assert [warning.code for warning in document.warnings] == ["STYLE_MAPPING_UNCERTAIN"]
